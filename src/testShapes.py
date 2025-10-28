import logging
import sys
import copy
import os
import time
import pathlib
from engine import Engine
import cv2 as cv
import numpy as np
from docx import Document
from docx.shared import Inches
from multiprocessing import Process, Pool, current_process
from progress.bar import IncrementalBar

DEBUG = True
PROCESS_COUNT = 15

logging.basicConfig(encoding="utf-8", level=logging.INFO,
                    format="%(levelname)s %(asctime)s %(message)s",
                    handlers=[logging.FileHandler("log.txt", ("w" if DEBUG else "w+")), logging.StreamHandler(sys.stdout)])

shapeQueue = []
colorQueue = []
pool = []


def loadTiles(filename, prefix, e):

    img = e.loadImage(filename)
    image_copy = img.copy() 
    imgheight=img.shape[0]
    imgwidth=img.shape[1]
    tileNames = []

    M = 500
    N = 500
    x1 = 0
    y1 = 0

    #if imgheight < M:
    #    logging.warning("Image height is smaller than a tile height! Undefined behaviour!")
    #
    #if imgwidth < N:
    #    logging.warning("Image width is smaller than a tile width! Undefined behaviour!")
     
    logging.info("Exporting image as tiles")

    for y in range(0, imgheight, M):
        tileNamesRow = []
        for x in range(0, imgwidth, N):
            if (imgheight - y) < M or (imgwidth - x) < N:
                name = 'saved_patches/'+str(prefix)+'_tile'+str(x)+'_'+str(y)+'.jpg'
                y1 = y+M
                x1 = x+N
                if (x1 >= imgwidth):
                    x1 = imgwidth-1

                if (y1 >= imgheight):
                    y1 = imgheight-1

                tiles = image_copy[y:y1, x:x1]
                e.exportImage(tiles, name)
                tileNamesRow.append(name)
                continue
                 
            y1 = y + M
            x1 = x + N

            name = 'saved_patches/'+str(prefix)+'_tile'+str(x)+'_'+str(y)+'.jpg'

            # check whether the patch width or height exceeds the image width or height
            if x1 >= imgwidth and y1 >= imgheight:
                x1 = imgwidth - 1
                y1 = imgheight - 1
                #Crop into patches of size MxN
                tiles = image_copy[y:y+M, x:x+N]
                #Save each patch into file directory
                e.exportImage(tiles, name)
                #cv.rectangle(img, (x, y), (x1, y1), (0, 255, 0), 1)
            elif y1 >= imgheight: # when patch height exceeds the image height
                y1 = imgheight - 1
                #Crop into patches of size MxN
                tiles = image_copy[y:y+M, x:x+N]
                #Save each patch into file directory
                e.exportImage(tiles, name)
                #cv.rectangle(img, (x, y), (x1, y1), (0, 255, 0), 1)
            elif x1 >= imgwidth: # when patch width exceeds the image width
                x1 = imgwidth - 1
                #Crop into patches of size MxN
                tiles = image_copy[y:y+M, x:x+N]
                #Save each patch into file directory
                e.exportImage(tiles, name)
                #cv.rectangle(img, (x, y), (x1, y1), (0, 255, 0), 1)
            else:
                #Crop into patches of size MxN
                tiles = image_copy[y:y+M, x:x+N]
                #Save each patch into file directory
                e.exportImage(tiles, name)
                #cv.rectangle(img, (x, y), (x1, y1), (0, 255, 0), 1)

            tileNamesRow.append(name)

        tileNames.append(tileNamesRow)
    return tileNames

def processTile(tile):
    tile1 = tile[0]
    tile2 = tile[1]
    prefix = str(os.getpid())
    colorQueue = []
    shapeQueue = []
    e = Engine()
    img1 = e.loadImage(tile1)
    img2 = e.loadImage(tile2)

    #os.remove(tile1)
    #os.remove(tile2)

    gray1 = e.grayscaleImage(img1)
    gray2 = e.grayscaleImage(img2)

    cont1 = e.findContours(gray1)
    cont2 = e.findContours(gray2)

    gray1 = cv.cvtColor(gray1, cv.COLOR_GRAY2BGR)
    gray2 = cv.cvtColor(gray2, cv.COLOR_GRAY2BGR)

    #e.drawContours(gray1, cont1)
    #e.drawContours(gray2, cont2)

    #e.viewImages(gray1, gray2)


    #for c1, c2 in zip(cont1, cont2):
    #    r = copy.deepcopy(img1)
    #    r2 = copy.deepcopy(img2)

    #    e.drawContours(r, [c1], color=(255,0,0))
    #    e.drawContours(r2, [c2], color=(255,0,0))

    #    e.viewImages(r, r2)

    folder, t1 = os.path.split(tile1)
    folder, t2 = os.path.split(tile2)
    t1 = pathlib.Path(t1).stem
    t2 = pathlib.Path(t2).stem

    #r = copy.deepcopy(img1)
    #r2 = copy.deepcopy(img2)

    # COLOR DIFF
    diff = e.findDiffColors(img1, img2, cont1, cont2)
    if diff is None:
        pass
        #logging.debug("No difference in colors found")
    else:
        #logging.debug("Found difference in colors!")
        for d in diff:

            r = copy.deepcopy(img1)
            r2 = copy.deepcopy(img2)

            e.drawContours(r, [d[0][0]], color=(255,0,0), thickness=e.calculateThickness(e.calculateOffset(d[0][0])))
            e.drawContours(r2, [d[0][1]], color=(255,0,0), thickness=e.calculateThickness(e.calculateOffset(d[0][0])))                    

            #doc.add_heading("Найдено различие в цвете")
            #p = doc.add_paragraph()
            #run = p.add_run()
            ##e.exportImage(e.cropToContour(r, d[0][0]), "tmp.jpg")
            #e.exportImage(r, str(prefix)+"_"+str(tile1)+"_tmp.jpg")
            #colorQueue.append()
            #run.add_picture(str(prefix)+"_tmp.jpg", width=Inches(2))
            #run = p.add_run()
            ##e.exportImage(e.cropToContour(r2, d[0][1]), "tmp.jpg")
            #e.exportImage(r2, str(prefix)+"_"+str(tile2)+"_tmp.jpg")
            #run.add_picture(str(prefix)+"_tmp.jpg", width=Inches(2))
            #os.remove(str(prefix)+"_tmp.jpg")

            e.exportImage(e.cropToContour(r, d[0][0]), os.path.join(folder, str(prefix)+"_"+str(t1)+"_tmp.jpg"))
            e.exportImage(e.cropToContour(r2, d[0][1]), os.path.join(folder, str(prefix)+"_"+str(t2)+"_tmp2.jpg"))

            colorQueue.append([os.path.join(folder, str(prefix)+"_"+str(t1)+"_tmp.jpg"), os.path.join(folder, str(prefix)+"_"+str(t2)+"_tmp2.jpg")])

    # SHAPE DIFF
    df = e.findDiffShapes(cont1, cont2)
    if df is None:
        pass
        #logging.debug("No difference in shapes found")
    else:
        #logging.debug("Found difference in shapes!")
        #e.drawContours(r, df[0], thickness=-1, color=(0,0,255)) # draw diff contours from other img
        #e.drawContours(r2, df[1], thickness=-1, color=(0,0,255))
        for c1, c2 in zip(df[0], df[1]):

            r = copy.deepcopy(img1)
            r2 = copy.deepcopy(img2)

            e.drawContours(r, [c1], thickness=e.calculateThickness(e.calculateOffset(c1)), color=(0,0,255))
            e.drawContours(r2, [c2], thickness=e.calculateThickness(e.calculateOffset(c2)), color=(0,0,255))

            #doc.add_heading("Различие в форме")
            #p = doc.add_paragraph()
            #run = p.add_run()
            #e.exportImage(e.cropToContour(r, c1), str(prefix)+"_"+str(tile1)+"_tmp.jpg") 
            #run.add_picture(str(prefix)+"_tmp.jpg", width=Inches(2))
            #run = p.add_run("    ")
            #e.exportImage(e.cropToContour(r2, c2), str(prefix)+"_"+str(tile2)+"_tmp.jpg")
            #run.add_picture(str(prefix)+"_tmp.jpg", width=Inches(2))
            #os.remove(str(prefix)+"_tmp.jpg")
            
            e.exportImage(e.cropToContour(r, c1), os.path.join(folder, str(prefix)+"_"+str(t1)+"_tmp3.jpg"))
            e.exportImage(e.cropToContour(r2, c2), os.path.join(folder, str(prefix)+"_"+str(t2)+"_tmp4.jpg"))

            shapeQueue.append([os.path.join(folder, str(prefix)+"_"+str(t1)+"_tmp3.jpg"), os.path.join(folder, str(prefix)+"_"+str(t2)+"_tmp4.jpg")])

    if df is not None:
        e.drawContours(img1, df[0], thickness=2, color=(0,0,255))
        e.drawContours(img2, df[1], thickness=2, color=(0,0,255))
    else:
        logging.debug("No difference in shapes found!")

    
    if diff is not None:
        for d in diff:
            e.drawContours(img1, [d[0][0]], color=(255,0,0), thickness=e.calculateThickness(e.calculateOffset(d[0][0])))
            e.drawContours(img2, [d[0][1]], color=(255,0,0), thickness=e.calculateThickness(e.calculateOffset(d[0][1])))
    else:
        logging.debug("No difference in colors found!")

    e.exportImage(img1, tile1)
    e.exportImage(img2, tile2)

    # export queues to the file
    with open("queues/"+str(prefix)+"_queue.txt", "a") as f:
        for c in colorQueue:
            f.write("color: "+c[0]+";"+c[1]+"\n")

        for s in shapeQueue:
            f.write("shape: "+s[0]+";"+s[1]+"\n")

def processTiles(firstTiles, secondTiles, e, doc: Document): 
    for p in os.listdir("queues"):
        os.remove(os.path.join("queues", p))
    logging.info("Starting " + str(PROCESS_COUNT) + " processes")
    for tileRow1, tileRow2 in zip(firstTiles, secondTiles):
        iter = zip(tileRow1, tileRow2)
        with Pool(PROCESS_COUNT) as p:
            p.map(processTile, iter)

    logging.info("Making summary")

    bar = IncrementalBar("Procesing", max=len(os.listdir("queues")))
    try:
        for ind, p in enumerate(os.listdir("queues")):
            with open(os.path.join("queues", p), "r") as f:
                sp = f.read().split("\n")
                for s in sp:
                    if s.startswith("shape: "):
                        try:
                            name = s[7:]
                            name1, name2 = name.split(';')
                            doc.add_heading("Различие в форме")
                            p = doc.add_paragraph()
                            run = p.add_run()
                            run.add_picture(name1, width=Inches(2))
                            run = p.add_run("    ")
                            run.add_picture(name2, width=Inches(2))
                        except KeyboardInterrupt:
                            doc.save("summary.docx")
                            raise KeyboardInterrupt
                        except Exception as e:
                            logging.error(f"Error occured while processing '{name}':\n'{str(e)}' and the file was {'found' if os.path.exists(name1) else 'not found'}")
                    elif s.startswith("color: "):
                        try:
                            name = s[7:]
                            name1, name2 = name.split(';')
                            doc.add_heading("Различие в цвете")
                            p = doc.add_paragraph()
                            run = p.add_run()
                            run.add_picture(name1, width=Inches(2))
                            run = p.add_run()
                            run.add_picture(name2, width=Inches(2))
                        except Exception as e:
                            logging.error(f"Error occured while processing '{name}':\n'{str(e)}' and the file was {'found' if os.path.exists(name1) else 'not found'}")
                f.close()
            bar.next()
    except KeyboardInterrupt:
        print("Summary interrupted")
    finally:
        bar.finish()
        doc.save("summary.docx")


    img = None
    for ind, tileRow in enumerate(firstTiles):
        row = None
        for ind2, tile in enumerate(tileRow):
            tileImg = e.loadImage(tile)
            #os.remove(tile)
            if ind2 == 0:
                row = tileImg
            else:
                row = np.concatenate((row, tileImg), axis=1)
        if row is None:
            break
        if ind == 0:
            img = row
        else:
            img = np.concatenate((img, row), axis=0)

    img2 = None
    for ind, tileRow in enumerate(secondTiles):
        row = None
        for ind2, tile in enumerate(tileRow):
            tileImg = e.loadImage(tile)
            os.remove(tile)
            if ind2 == 0:
                row = tileImg
            else:
                row = np.concatenate((row, tileImg), axis=1)
        if row is None:
            break
        if ind == 0:
            img2 = row
        else:
            img2 = np.concatenate((img2, row), axis=0)

    return img, img2

def processPair(first, second):

    #cleanup
    if os.path.exists("saved_patches/"):
        for p in os.listdir("saved_patches"):
            os.remove(os.path.join("saved_patches", p))
    else:
        os.mkdir("saved_patches")

    if os.path.exists("queues/"):
        for p in os.listdir("queues"):
            os.remove(os.path.join("queues", p))
    else:
        os.mkdir("queues")

    logging.info("Processing pair")
    e = Engine()
    doc = Document()

    firstTiles = loadTiles(first, "first", e)
    secondTiles = loadTiles(second, "second", e)

    start = time.time()
    img, img2 = processTiles(firstTiles, secondTiles, e, doc)
    if img is None:
        logging.error("First result image is None! Exporting first image...")
        e.exportImage(e.loadImage(first), "result1.jpg")
    else:
        e.exportImage(img, "result1.png")

    if img2 is None:
        logging.error("Second result image is None! Exporting second image...")
        e.exportImage(e.loadImage(second), "result2.jpg")
    else:
        e.exportImage(img2, "result2.png")
    end = time.time()
    logging.info(f"Process took {round(end-start, 2)}s")

if __name__ == "__main__":
    processPair("map1.png", "map2.png")
    #processPair("test1.png", "test2.png")
    #e = Engine()
    #img1 = e.loadImage("test_c1.jpg")
    #img2 = e.loadImage("test_c2.jpg")
